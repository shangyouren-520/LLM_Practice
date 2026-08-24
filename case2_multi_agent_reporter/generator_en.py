"""
=============================================================================
Case 2: Multi-Agent Multimodal Financial Research Report System - ResearchManager
(Workflow Orchestration and Word Assembly Module)
-----------------------------------------------------------------------------
[Teaching Objective]:
This module implements the Research Manager and multimodal assembly workflow
introduced on Slides 48-51. It uses a classic, intuitive Manager-Workers
architecture to orchestrate four major stages:
1. Stage 1: Plan -> invoke PlannerAgent to generate the report outline;
2. Stage 2: Data -> invoke DataEngineerAgent to obtain the market-data package and chart;
3. Stage 3: Analyze -> repeatedly invoke AnalystAgent to write section-level analysis;
4. Stage 4: Assemble -> use python-docx to combine text, trend charts, and tabular data
   into a standardized multimodal financial research report.
=============================================================================
"""

import os
import logging
from typing import Dict, List, Any

# python-docx is used to generate standardized Word documents.
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

from llm_factory_en import LLMFactory
from agents_en import PlannerAgent, DataEngineerAgent, AnalystAgent

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")


class ResearchManager:
    """
    Research Manager acting as the project-level orchestrator.

    It coordinates agents with different responsibilities, ensures each task is
    executed at the appropriate stage, and ultimately delivers a multimodal
    financial research report.
    """

    def __init__(self, llm_factory: LLMFactory = None):
        self.llm_factory = llm_factory or LLMFactory()
        self.planner = PlannerAgent(llm_factory=self.llm_factory)
        self.data_engineer = DataEngineerAgent()
        self.analyst = AnalystAgent(llm_factory=self.llm_factory)

    def generate_report(self, stock_name: str = "CATL", stock_code: str = "300750") -> Dict[str, str]:
        """
        Launch the fully automated multi-agent research-report workflow with one call
        (corresponds to Slide 41).
        """
        print("\n" + "=" * 60)
        logging.info(
            f"🚀 Starting the autonomous research-report workflow for "
            f"[{stock_name} ({stock_code})]..."
        )
        print("=" * 60)

        # ---------------------------------------------------------------------
        # Stage 1: Outline planning (corresponds to Stage 1 on Slides 48/49)
        # ---------------------------------------------------------------------
        logging.info("👉 [Stage 1/4] Launching PlannerAgent to define the report outline...")
        sections = self.planner.run({
            "input": f"Plan the core section structure for the {stock_name} research report."
        })

        # ---------------------------------------------------------------------
        # Stage 2: Data extraction and charting (corresponds to Stage 2 on Slides 48/49)
        # ---------------------------------------------------------------------
        logging.info(
            "👉 [Stage 2/4] Launching DataEngineerAgent to retrieve market data "
            "and generate a high-resolution trend chart..."
        )
        data_packet = self.data_engineer.run(stock_code=stock_code, stock_name=stock_name)

        # ---------------------------------------------------------------------
        # Stage 3: Section-level analysis (corresponds to Stage 3 on Slides 48/49)
        # ---------------------------------------------------------------------
        logging.info(
            "👉 [Stage 3/4] Launching AnalystAgent to write in-depth analysis "
            "for each report section..."
        )
        word_sections = []
        for sec_title in sections:
            logging.info(f"✍️ Writing section: \"{sec_title}\"...")
            content = self.analyst.run({
                "input": (
                    f"Section title: {sec_title}\n"
                    f"Data context: {data_packet['data_context']}"
                )
            })
            word_sections.append({"title": sec_title, "content": content})

        # ---------------------------------------------------------------------
        # Stage 4: Multimodal assembly and formatting (Slides 48/50, Stage 4)
        # ---------------------------------------------------------------------
        logging.info(
            "👉 [Stage 4/4] Assembling the final Word research report and Markdown file..."
        )
        docx_path = self._save_docx(stock_name, stock_code, data_packet, word_sections)
        md_path = self._save_markdown(stock_name, stock_code, data_packet, word_sections)

        print("\n" + "=" * 60)
        logging.info("🎉 Multi-agent research report generation completed successfully!")
        logging.info(f"📄 DOCX deliverable: {os.path.abspath(docx_path)}")
        logging.info(f"📝 Markdown deliverable: {os.path.abspath(md_path)}")
        print("=" * 60 + "\n")

        return {"docx": docx_path, "markdown": md_path}

    def _save_docx(
        self,
        stock_name: str,
        stock_code: str,
        data_packet: Dict[str, Any],
        word_sections: List[Dict[str, str]]
    ) -> str:
        """
        Assemble multimodal content into a professionally formatted Word research
        report using python-docx (corresponds to Slide 50).
        """
        output_dir = "output_reports"
        os.makedirs(output_dir, exist_ok=True)
        docx_file = os.path.join(
            output_dir,
            f"{stock_name}_{stock_code}_In-Depth_Research_Report.docx"
        )

        if Document is None:
            logging.warning("⚠️ python-docx is not installed; skipping Word document generation.")
            return ""

        doc = Document()

        # 1. Configure the document font (Slide 50 standard: Microsoft YaHei).
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        # 2. Insert the main title.
        title_para = doc.add_heading(
            f"{stock_name} ({stock_code}) In-Depth Research Report",
            level=0
        )
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(
            f"Report Date: {data_packet['df']['Date'].iloc[-1]}  |  "
            f"Subject: {stock_name}  |  "
            "Industry: New Energy and Advanced Manufacturing\n"
        )

        # 3. Insert the primary trend chart (visual element).
        doc.add_heading("1. Core Price Trend", level=1)
        if data_packet.get("chart_path") and os.path.exists(data_packet["chart_path"]):
            doc.add_picture(data_packet["chart_path"], width=Inches(6.0))
            last_pic = doc.paragraphs[-1]
            last_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Trend Summary: {data_packet['data_context']}")

        # 4. Insert the analysis sections (text elements).
        for i, sec in enumerate(word_sections, start=2):
            doc.add_heading(f"{sec['title']}", level=1)
            p = doc.add_paragraph(sec["content"])
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.space_after = Pt(8)

        # 5. Insert a recent-trading-data table (structured data; Slides 50/51).
        doc.add_heading("Appendix: Selected Recent Trading Data", level=1)
        df_recent = data_packet["df"].tail(5)  # Use the most recent five trading days.

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # Populate the table header.
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Date'
        hdr_cells[1].text = 'Close (CNY)'
        hdr_cells[2].text = 'High (CNY)'
        hdr_cells[3].text = 'Volume (lots)'

        # Populate data rows.
        for _, row in df_recent.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['Date'])
            row_cells[1].text = f"{row['Close']:.2f}"
            row_cells[2].text = f"{row['High']:.2f}"
            row_cells[3].text = f"{int(row['Volume']):,}"

        doc.save(docx_file)
        return docx_file

    def _save_markdown(
        self,
        stock_name: str,
        stock_code: str,
        data_packet: Dict[str, Any],
        word_sections: List[Dict[str, str]]
    ) -> str:
        """Save a Markdown copy for quick preview on GitHub or the web."""
        output_dir = "output_reports"
        os.makedirs(output_dir, exist_ok=True)
        md_file = os.path.join(
            output_dir,
            f"{stock_name}_{stock_code}_In-Depth_Research_Report.md"
        )

        md_content = f"# {stock_name} ({stock_code}) In-Depth Research Report\n\n"
        md_content += f"> **Data Baseline**: {data_packet['data_context']}\n\n"
        md_content += "## 1. Core Market Trend\n\n"
        if data_packet.get("chart_path"):
            md_content += f"![Trend Chart](../{data_packet['chart_path']})\n\n"

        for sec in word_sections:
            md_content += f"## {sec['title']}\n\n{sec['content']}\n\n"

        md_content += "## Appendix: Trading Data for the Most Recent 5 Days\n\n"
        md_content += (
            "| Date | Close (CNY) | High (CNY) | Volume (lots) |\n"
            "| :--- | :--- | :--- | :--- |\n"
        )
        df_recent = data_packet["df"].tail(5)
        for _, row in df_recent.iterrows():
            md_content += (
                f"| {row['Date']} | {row['Close']:.2f} | "
                f"{row['High']:.2f} | {int(row['Volume']):,} |\n"
            )

        with open(md_file, "w", encoding="utf-8") as f:
            f.write(md_content)

        return md_file


# -----------------------------------------------------------------------------
# Standalone CLI entry point (corresponds to the one-line call on Slide 41)
# -----------------------------------------------------------------------------
if __name__ == "__main__":
    manager = ResearchManager()
    manager.generate_report(stock_name="CATL", stock_code="300750")
